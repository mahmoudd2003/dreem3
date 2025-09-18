# utils/exporters.py
from typing import Dict, Any, List
from io import BytesIO
from docx import Document
from docx.shared import Pt
import json

def to_markdown(article_markdown: str) -> str:
    """يرجع محتوى الماركداون كما هو (للتنزيل كـ .md)."""
    return article_markdown or ""

def to_docx_bytes(article_markdown: str, meta_title: str = "") -> bytes:
    """
    يحوّل الماركداون لنص عادي داخل DOCX (تحويل مبسّط بدون تنسيق Markdown متقدم).
    إن أردت تحويل Markdown كامل لاحقًا، يمكن دمج محوّل مخصص.
    """
    doc = Document()

    if meta_title:
        h = doc.add_heading(meta_title, level=1)
        for r in h.runs:
            r.font.name = 'Arial'
            r.font.size = Pt(16)

    for line in (article_markdown or "").splitlines():
        p = doc.add_paragraph(line)
        for r in p.runs:
            r.font.name = 'Arial'
            r.font.size = Pt(11)

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def to_json_bytes(
    article_markdown: str,
    meta: Dict[str, Any],
    keyword: str,
    related_keywords: List[str],
) -> bytes:
    payload = {
        "keyword": keyword,
        "related_keywords": related_keywords,
        "meta": meta or {},
        "article_markdown": article_markdown or "",
    }
    return json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")
